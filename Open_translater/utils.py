import os
import uuid
from googletrans import Translator
from docx import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from fpdf import FPDF

translator = Translator()

def handle_uploaded_file(filepath, source_lang, target_lang, output_folder):

    translator = Translator()

    filename = os.path.basename(filepath)
    name, ext = os.path.splitext(filename)
    ext = ext.lower()

    if ext == '.docx':
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == '.pdf':
        reader = PdfReader(filepath)
        text = "\n".join([page.extract_text() or '' for page in reader.pages])
    else:
        raise ValueError("지원하지 않는 파일 형식입니다.")

    translated = translator.translate(text, src=source_lang, dest=target_lang).text

    safe_name = uuid.uuid4().hex

    if ext == '.docx':
        output_filename = f"{safe_name}_translated.docx"
        output_path = os.path.join(os.path.abspath(output_folder), output_filename)

        new_doc = DocxDocument()
        for line in translated.split("\n"):
            new_doc.add_paragraph(line)
        new_doc.save(output_path)

    elif ext == '.pdf':
        output_filename = f"{safe_name}_translated.pdf"
        output_path = os.path.join(os.path.abspath(output_folder), output_filename)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # ✅ TTF 폰트 등록 (예: 나눔고딕, 맑은 고딕 등)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        font_path = os.path.join(base_dir, "NanumGothic.ttf")
        pdf.add_font("Nanum", "", font_path, uni=True)
        pdf.set_font("Nanum", size=12)

        for line in translated.split("\n"):
            pdf.multi_cell(0, 10, line)
        pdf.output(output_path)

    return translated, output_path